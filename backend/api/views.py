import subprocess
import os
from docx import Document
import tempfile
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse,FileResponse
from openai import OpenAI
import json
import cv2
from pdf2image import convert_from_path
import paddlehub as hub
import logging
import re


@csrf_exempt
def MaterialContract(request):
    if request.method == 'POST' and request.FILES.get('file'):
        upload = request.FILES['file']
        ext = os.path.splitext(upload.name)[-1].lower()

        # 保存文件到临时路径
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext ,dir=r'C:\Python project\ContractAnalysis\backend\uploads') as tmp:
            for chunk in upload.chunks():
                tmp.write(chunk)
            file_path = tmp.name

        if ext == '.docx':
            text = extract_docx_text(file_path)
        elif ext == '.doc':
            text = convert_doc_to_docx_and_extract(file_path)
        elif ext == '.pdf':
            content = pdf_ocr_extract_text(file_path)
            # 构造 prompt 提交到 DeepSeek
            prompt = f"""
                        这是一份特殊格式的购销合同：{content},它是一个列表，里面有很多个字典，每个字典的键是合同原文，值是这段原文在合同中的位置坐标，
                        请你作为合同审核专家，从以下9个方面分析购销合同内容是否符合公司合规要求：
                        1.供方资质审核：是否存在失信情况，是否有重大负债（100万元以上）；
                        2.合同编号中日期与签订时间一致；
                        3.产品名称前后一致，合同总金额计算无误，大小写一致；
                        4.交货时间和地点清晰明确；
                        5.结算方式为承兑付款，付款时间清晰明确；
                        6.违约责任包含逾期到货违约金和逾期到货合同解除权；
                        7.争议解决的诉讼法院为需方所在地或原告方所在法院；
                        8.签订时间已明确，无空白条款。
                        9.该合同是否与《民法典》及公司规章制度相冲突。
                        对于每一个点，你需要返回4个字段：
                        "status": "符合 /风险 /不符"，若与要求合规则返回"符合"；若与要求基本符合但有别的风险则返回"风险"；若与要求不一致或无法查证则返回"不符"。
                        "original_text": "返回该点在合同中的原文，且必须和原文一字不差，若合同原文中没有找到这点相关的信息，则输出原文未提及",
                        "detail": "对于这点的分析详情与风险提示",
                        “highlights": "合同原文对应的位置坐标,是一个list，若对应原文多处则包含多个坐标”,
                        你需要严格按照这个格式返回数据，最后将这个九个点的分析结果按顺序以JSON格式返回，示例输出格式如下：
                        [
                          {{
                            "status": "不符",
                            "original_text": "原文未提及",
                            "detail": "合同中未涉及供方（景津装备股份有限公司）的失信记录或负债情况。需通过外部渠道（如国家企业信用信息公示系统、天眼查等）独立核实其是否存在失信、重大负债（>100万元）等问题。",
                            "highlights": [{{'page': 1,'x': 100,'y': 200,'width': 200,'height': 50}}]
                          }},
                          {{
                            "status": "符合",
                            "original_text": "合同编号：JAJJ20250507504 签订日期：2025年5月7日"
                            "detail": "合同编号中的日期为：20250507（2025年5月7日），与签订日期2025年5月7日一致。",
                            "highlights": [{{'page': 1,'x': 100,'y': 200,'width': 200,'height': 50}},{{'page': 1,'x': 100,'y': 200,'width': 200,'height': 50}}]
                          }},

                          (......中间6个点以此类推)

                          {{
                            "status": "风险",
                            "original_text": '如果供方没有按照合同规定的时间交货，供方应支付需方违约金，违约金应按每迟
            交一天，按迟期交货部分价款的0.5％计收。逾期一周以上的，需方有权单方面解除本合同，供方应退还需方支付的全部款项，并按合同总价的20%支付违约金。'

                            "detail": '《民法典》合规性分析：
            违约金比例（日0.5% ≈ 年182.5%）：
            ⚠️ 可能过高（《民法典》第585条：违约金不得超过实际损失30%），但需方可通过诉讼请求调减。

            其他条款（交货、付款、争议解决）无冲突。

            公司制度冲突：
            🔍 需内部比对（合同未体现需方内部规则）。'
                          }},
                          "highlights": [{{'page': 1,'x': 100,'y': 200,'width': 200,'height': 50}}]
                        ]

                        而对于第八点空白条款的判定，你需要结合以下合同模板进行审核：
                        购销合同
                                                                   合同编号：
            供方：                                                               签订方式： 
            需方：江西晶安          有限公司                                     签订时间：
            根据合同法等相关法律法规，供需双方经协商一致，就需方采购                  相关事宜，签订本合同。
            一、产品名称、型号规格、数量、金额等。
            产品名称	规格	计量单位	数 量	单价(元)	总价（元）	备注

            合计						
            合计人民币：                         （含17%增值税及运费）
            二、质量标准：                                                                               
            三、供货时间及交货地点：                                                                   。
            四、付款方式：                                                                             。
            五、质量保证与服务条款：                                                                   。
            六、验收标准、方法及提出异议：                                                             。  
            七、争议解决：若合同履行中，发生争议，双方应协商解决，协商不成时，由需方所在地法院依法裁决 。
            八、其他约定事项：
            1、供方必须按约如期交货，如供方不能按约交货，需提前7天通知需方，如未通知而延期交货给需方造成的所有损失，由供方承担。
            2、如供方逾期交货超过7天，则供方应按合同总金额的20%向支付违约金，同时需方有权选择是否继续执行合同。如供方逾期交货造成需方从市场上高价购入该货物时，则供方应按需方所购货物包含运费在内实际差价的1.2倍给予赔偿，因此造成需方其他损失的，供方应予以赔偿。
            3、若供方所提供的设备不符合质量要求，应负责更换或退货，并承担相应责任与费用，若在需方通知后的合理时间内，供方不予更换或退货，或者更换后质量仍不符合要求，则需方有权解除本合同，供方应退还需方支付的全部款项，并按合同金额的20%向需方承担违约责任。
            4、本合同一式两份，甲乙双方各执一份，经双方签字盖章后生效，传真件、扫描件与原件具有同等效力。
            供方	需方
            单位名称：
            单位地址：
            法定代表人：
            经办人：
            电话：                
            传真：
            开户银行：
            帐号：	单位名称：江西晶安       有限公司
            单位地址：南昌市安义县万埠镇八宝路37号
            经办人：
            法定代表人：
            电话：0791-83432589                  
            传真：0791-83432197 
            开户银行：
            帐号：

            要求合同中的条款要与上述模板对应，不得缺少或者漏填。
                        请你将分析结果按格式返回。
                        """
            # 设置 DeepSeek 客户端
            client = OpenAI(api_key="sk-ca6282f72dcd42798d469ea66ffc02a2", base_url="https://api.deepseek.com")

            response = client.chat.completions.create(
                model="deepseek-reasoner",
                messages=[
                    {"role": "system", "content": "你是一个专业的企业合同审核助手"},
                    {"role": "user", "content": prompt},
                ],
                stream=False
            )

            # print(response.choices[0].message.content)
            result_text = response.choices[0].message.content

            # 尝试解析返回为 JSON，如果失败就直接返回原始字符串
            result_text = parse_json_with_code_blocks(result_text)
            print(result_text)
            os.remove(file_path)
            return JsonResponse(result_text, safe=False)
        else:
            return JsonResponse({'error': 'Unsupported file type. Please upload a .doc or .docx file.'}, status=400)
        # if text:
        #     print("从docx文件提取的文本:")
        #     print(text)  # 打印前200个字符

        # 构造 prompt 提交到 DeepSeek
        prompt = f"""
            请你作为合同审核专家，从以下9个方面分析购销合同内容是否符合公司合规要求：
            1.供方资质审核：是否存在失信情况，是否有重大负债（100万元以上）；
            2.合同编号中日期与签订时间一致；
            3.产品名称前后一致，合同总金额计算无误，大小写一致；
            4.交货时间和地点清晰明确；
            5.结算方式为承兑付款，付款时间清晰明确；
            6.违约责任包含逾期到货违约金和逾期到货合同解除权；
            7.争议解决的诉讼法院为需方所在地或原告方所在法院；
            8.签订时间已明确，无空白条款。
            9.该合同是否与《民法典》及公司规章制度相冲突。
            对于每一个点，你需要返回三个字段：
            "status": "符合 /风险 /不符"，若与要求合规则返回"符合"；若与要求基本符合但有别的风险则返回"风险"；若与要求不一致或无法查证则返回"不符"。
            "original_text": "返回该点在合同中的原文，且必须和原文一字不差，若合同原文中没有找到这点相关的信息，则输出原文未提及",
            "detail": "对于这点的分析详情与风险提示",
            你需要严格按照这个格式返回数据，最后将这个九个点的分析结果按顺序以JSON格式返回，示例输出格式如下：
            [
              {{
                "status": "不符",
                "original_text": "原文未提及"
                "detail": "合同中未涉及供方（景津装备股份有限公司）的失信记录或负债情况。需通过外部渠道（如国家企业信用信息公示系统、天眼查等）独立核实其是否存在失信、重大负债（>100万元）等问题。",
              }},
              {{
                "status": "符合",
                "original_text": "合同编号：JAJJ20250507504 签订日期：2025年5月7日"
                "detail": "合同编号中的日期为：20250507（2025年5月7日），与签订日期2025年5月7日一致。",
              }},
              
              (......中间6个点以此类推)
              
              {{
                "status": "风险",
                "original_text": 如果供方没有按照合同规定的时间交货，供方应支付需方违约金，违约金应按每迟
交一天，按迟期交货部分价款的0.5％计收。逾期一周以上的，需方有权单方面解除本合同，供方应退还需方支付的全部款项，并按合同总价的20%支付违约金。
            
                "detail": 《民法典》合规性分析：
违约金比例（日0.5% ≈ 年182.5%）：
⚠️ 可能过高（《民法典》第585条：违约金不得超过实际损失30%），但需方可通过诉讼请求调减。

其他条款（交货、付款、争议解决）无冲突。

公司制度冲突：
🔍 需内部比对（合同未体现需方内部规则）。
              }},
            ]
            
            而对于第八点空白条款的判定，你需要结合以下合同模板进行审核：
            购销合同
                                                       合同编号：
供方：                                                               签订方式： 
需方：江西晶安          有限公司                                     签订时间：
根据合同法等相关法律法规，供需双方经协商一致，就需方采购                  相关事宜，签订本合同。
一、产品名称、型号规格、数量、金额等。
产品名称	规格	计量单位	数 量	单价(元)	总价（元）	备注
						
合计						
合计人民币：                         （含17%增值税及运费）
二、质量标准：                                                                               
三、供货时间及交货地点：                                                                   。
四、付款方式：                                                                             。
五、质量保证与服务条款：                                                                   。
六、验收标准、方法及提出异议：                                                             。  
七、争议解决：若合同履行中，发生争议，双方应协商解决，协商不成时，由需方所在地法院依法裁决 。
八、其他约定事项：
1、供方必须按约如期交货，如供方不能按约交货，需提前7天通知需方，如未通知而延期交货给需方造成的所有损失，由供方承担。
2、如供方逾期交货超过7天，则供方应按合同总金额的20%向支付违约金，同时需方有权选择是否继续执行合同。如供方逾期交货造成需方从市场上高价购入该货物时，则供方应按需方所购货物包含运费在内实际差价的1.2倍给予赔偿，因此造成需方其他损失的，供方应予以赔偿。
3、若供方所提供的设备不符合质量要求，应负责更换或退货，并承担相应责任与费用，若在需方通知后的合理时间内，供方不予更换或退货，或者更换后质量仍不符合要求，则需方有权解除本合同，供方应退还需方支付的全部款项，并按合同金额的20%向需方承担违约责任。
4、本合同一式两份，甲乙双方各执一份，经双方签字盖章后生效，传真件、扫描件与原件具有同等效力。
供方	需方
单位名称：
单位地址：
法定代表人：
经办人：
电话：                
传真：
开户银行：
帐号：	单位名称：江西晶安       有限公司
单位地址：南昌市安义县万埠镇八宝路37号
经办人：
法定代表人：
电话：0791-83432589                  
传真：0791-83432197 
开户银行：
帐号：

要求合同中的条款要与上述模板对应，不得缺少或者漏填。

            需审核的合同原文如下：
            {text}

            请你将分析结果按格式返回。
            """
        # 设置 DeepSeek 客户端
        client = OpenAI(api_key="sk-ca6282f72dcd42798d469ea66ffc02a2", base_url="https://api.deepseek.com")

        response = client.chat.completions.create(
            model="deepseek-reasoner",
            messages=[
                {"role": "system", "content": "你是一个专业的企业合同审核助手"},
                {"role": "user", "content": prompt},
            ],
            stream=False
        )

        # print(response.choices[0].message.content)
        result_text = response.choices[0].message.content

        # 尝试解析返回为 JSON，如果失败就直接返回原始字符串
        result_text = parse_json_with_code_blocks(result_text)
        result_text = annotate_items_with_positions(file_path, result_text)
        os.remove(file_path)
        #print(result_text)
        return JsonResponse(result_text, safe=False)

    else:
        return JsonResponse({'error': '请通过 POST 上传 doc 或 docx 文件'}, status=400)

# 提取docx的方法
def extract_docx_text(file_path):
    """使用python-docx从docx文件中提取纯文本内容"""
    try:
        doc = Document(file_path)
        full_text = []

        # 遍历文档中的每个段落
        for para in doc.paragraphs:
            full_text.append(para.text)

        # 遍历文档中的每个表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        # 将所有文本连接成一个字符串
        return '\n'.join(full_text)
    except Exception as e:
        print(f"提取docx文本失败: {e}")
        return None

# 提取doc的方法
def convert_doc_to_docx_and_extract(file_path):
    """使用 LibreOffice 将 .doc 转换为 .docx 并提取文本"""
    safe_file_path = os.path.abspath(file_path)
    if not os.path.exists(safe_file_path):
        print("文件不存在:", safe_file_path)
        return None

    # 设置输出路径（与原路径相同，添加后缀）
    output_dir = os.path.dirname(safe_file_path)
    base_name = os.path.splitext(os.path.basename(safe_file_path))[0]
    docx_path = os.path.join(output_dir, base_name + ".docx")

    try:
        # 调用 LibreOffice 进行转换
        subprocess.run([
            r"C:\LibreOffice\program\soffice.exe",
            "--convert-to", "docx",
            "--outdir", output_dir,
            safe_file_path
        ], check=True)

        # 检查转换结果
        if not os.path.exists(docx_path):
            print("转换失败，找不到生成的 .docx 文件")
            return None

        # 提取文本
        text = extract_docx_text(docx_path)

        # 删除中间文件
        os.remove(docx_path)

        return text
    except subprocess.CalledProcessError as e:
        print(f"LibreOffice 转换失败: {e}")
        return None
    except Exception as e:
        print(f"转换过程中发生错误: {e}")
        return None


def pdf_ocr_extract_text(pdf_path, output_dir=r"C:\Python project\ContractAnalysis\outimg"):
    os.makedirs(output_dir, exist_ok=True)

    # 1. PDF转JPG
    images = convert_from_path(pdf_path, dpi=300)
    jpg_paths = []
    for i, image in enumerate(images):
        image_path = os.path.join(output_dir, f"page_{i + 1}.jpg")
        image.save(image_path)
        jpg_paths.append(image_path)

    # 2. 初始化PaddleHub OCR模块
    ocr = hub.Module(name="ch_pp-ocrv3", enable_mkldnn=True)

    # 3. 读取图片并批量识别
    img_list = [cv2.imread(p) for p in jpg_paths]
    results = ocr.recognize_text(images=img_list)

    # 4. 提取文字内容，合并为字符串
    content = []
    for i,page_result in enumerate(results):
        img = img_list[i]
        img_h, img_w = img.shape[:2]  # 原图宽高
        for e in page_result['data']:
            point = convert_quad_to_rect(e['text_box_position'],i)
            point['x'] /= img_w
            point['y'] /= img_h
            point['width'] /= img_w
            point['height'] /= img_h    # 归一化
            content.append({e['text']:point})
    return content

# 输出清洗
def parse_json_with_code_blocks(json_text):
    """
    解析带Markdown代码块标记的JSON字符串
    支持处理 ```json 和 ``` 包裹的JSON数据
    """
    try:
        # 去除开头的 ```json 或 ```
        if json_text.startswith("```json"):
            json_text = json_text[len("```json"):]
        elif json_text.startswith("```"):
            json_text = json_text[len("```"):]

        # 去除末尾的 ```
        if json_text.endswith("```"):
            json_text = json_text[:-len("```")]

        # 去除前后可能存在的空白字符
        json_text = json_text.strip()

        # 解析JSON
        data = json.loads(json_text)
        return data
    except json.JSONDecodeError as e:
        print(f"JSON解析失败: {e}")
        print(f"待解析文本: {json_text[:100]}...")  # 打印前100个字符用于调试
        return None
    except Exception as e:
        print(f"处理过程中发生错误: {e}")
        return None


# ocr文字坐标转换
def convert_quad_to_rect(quad,page):
    if not isinstance(quad, list) or len(quad) != 4:
        raise ValueError("输入必须是四个点的列表，例如 [[x1,y1],[x2,y2],[x3,y3],[x4,y4]]")

    xs = [pt[0] for pt in quad]
    ys = [pt[1] for pt in quad]

    x = min(xs)
    y = min(ys)
    width = max(xs) - x
    height = max(ys) - y

    return {
        "page": page+1,
        "x": x,
        "y": y-height,
        "width": width,
        "height": height
    }


# word转pdf
import uuid
from django.views.decorators.csrf import csrf_exempt

@csrf_exempt
def ConvertToPdf(request):
    if request.method != "POST":
        return JsonResponse({"error": "只支持 POST 请求"}, status=405)

    uploaded_file = request.FILES.get("file")
    if not uploaded_file:
        return JsonResponse({"error": "未上传文件"}, status=400)

    if not uploaded_file.name.lower().endswith(('.doc', '.docx')):
        return JsonResponse({"error": "仅支持 .doc 或 .docx 文件"}, status=400)

    TEMP_ROOT = r'C:\Python project\ContractAnalysis\PDF'
    shutil.rmtree(TEMP_ROOT)
    os.makedirs(TEMP_ROOT, exist_ok=True)

    # 清理文件名中的非法字符
    def safe_filename(filename):
        return re.sub(r'[^\w\u4e00-\u9fa5.-]', '_', filename)

    tmpdir = os.path.join(TEMP_ROOT, str(uuid.uuid4()))
    os.makedirs(tmpdir, exist_ok=True)

    try:
        safe_name = safe_filename(uploaded_file.name)
        input_path = os.path.join(tmpdir, safe_name)

        # 保存上传的文件
        with open(input_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)

        soffice_path = r"C:\LibreOffice\program\soffice.exe"
        subprocess.run([
            soffice_path,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", tmpdir,
            input_path
        ], check=True)

        pdf_path = os.path.join(tmpdir, "temp.pdf")
        auto_generated_pdf = os.path.join(tmpdir, os.path.splitext(safe_name)[0] + ".pdf")
        if os.path.exists(auto_generated_pdf):
            os.rename(auto_generated_pdf, pdf_path)

        if not os.path.exists(pdf_path):
            return JsonResponse({"error": "转换失败，未生成 PDF"}, status=500)

        response = FileResponse(open(pdf_path, 'rb'), content_type='application/pdf', as_attachment=True, filename="temp.pdf")
        return response

    except subprocess.CalledProcessError as e:
        return JsonResponse({"error": f"LibreOffice 转换失败: {str(e)}"}, status=500)
    except Exception as e:
        return JsonResponse({"error": f"服务器异常: {str(e)}"}, status=500)




import fitz  # PyMuPDF
import shutil
from rapidfuzz import fuzz

def extract_word_positions_from_docx(docx_path, search_text):
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"文件不存在: {docx_path}")

    BASE_TMP_DIR = r'C:\Python project\ContractAnalysis\cache'
    os.makedirs(BASE_TMP_DIR, exist_ok=True)

    task_dir = os.path.join(BASE_TMP_DIR, str(uuid.uuid4()))
    os.makedirs(task_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    tmp_docx_path = os.path.join(task_dir, base_name + ".docx")
    shutil.copy2(docx_path, tmp_docx_path)

    results = []

    try:
        # 1. 转换 DOCX -> PDF
        subprocess.run([
            r"C:\LibreOffice\program\soffice.exe",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", task_dir,
            tmp_docx_path
        ], check=True, capture_output=True, text=True)

        pdf_path = os.path.join(task_dir, base_name + ".pdf")
        if not os.path.isfile(pdf_path):
            raise FileNotFoundError(f"PDF 转换失败，未找到输出文件：{pdf_path}")

        # 2. 分割关键词
        sub_texts = [t.strip() for t in search_text.split(' ') if t.strip()]

        with fitz.open(pdf_path) as doc:
            for page_number, page in enumerate(doc, start=1):
                page_width = page.rect.width
                page_height = page.rect.height
                # 提取段落文本块（含坐标）
                blocks = page.get_text("blocks")
                for block in blocks:
                    x0, y0, x1, y1, text = block[:5]
                    if not text.strip():
                        continue
                    for sub_text in sub_texts:
                        similarity = fuzz.partial_ratio(sub_text, text)
                        if similarity >= 95:
                            converted_y = y0-(y1 - y0)  # 左上角基准
                            results.append({
                                "page": page_number,
                                "x": x0/page_width,
                                "y": converted_y/page_height,
                                "width": (x1 - x0)/page_width,
                                "height": (y1 - y0)/page_height,
                            })


        return results

    except Exception as e:
        raise RuntimeError(f"处理失败：{e}")
    finally:
        try:
            shutil.rmtree(task_dir)
        except Exception as e:
            print(f"删除临时目录失败：{e}")


def annotate_items_with_positions(file_path, items):
    from copy import deepcopy
    results = deepcopy(items)

    for item in results:
        search_text = item.get("original_text", "").strip()
        if not search_text:
            item["highlights"] = []
            continue

        try:
            positions = extract_word_positions_from_docx(file_path, search_text)
        except Exception as e:
            # 如果转换或搜索失败，返回空结果
            print(f"处理失败：{search_text}，错误：{e}".encode("gbk", errors="ignore").decode("gbk"))
            positions = []

        item["highlights"] = positions

    return results




# views.py
from django.views.decorators.csrf import csrf_exempt
from django.utils import timezone
from .models import ContractAnalysis
import base64
from django.core.files.base import ContentFile

@csrf_exempt
def save_analysis(request):
    if request.method != 'POST':
        return JsonResponse({'error': '仅支持 POST 请求'}, status=405)

    try:
        if not request.body:
            return JsonResponse({'error': '空请求体'}, status=400)

        body_unicode = request.body.decode('utf-8')
        data = json.loads(body_unicode)

        file_name = data.get('file_name')
        audit_results = data.get('audit_results')
        pdf_base64 = data.get('pdf_base64')

        if not all([file_name, audit_results, pdf_base64]):
            return JsonResponse({'error': '参数不完整'}, status=400)

        pdf_data = base64.b64decode(pdf_base64)
        pdf_file = ContentFile(pdf_data, name=file_name.replace(' ', '_'))

        instance = ContractAnalysis.objects.create(
            file_name=file_name,
            pdf_file=pdf_file,
            audit_results=json.loads(audit_results),
            analysis_time=timezone.now()
        )

        return JsonResponse({'message': '保存成功', 'id': instance.id})

    except json.JSONDecodeError as e:
        return JsonResponse({'error': f'JSON 解析错误: {str(e)}'}, status=400)
    except Exception as e:
        return JsonResponse({'error': f'服务异常: {str(e)}'}, status=500)


# 返回历史数据api
def get_analysis_history(request):
    if request.method == 'GET':
        history = ContractAnalysis.objects.all().order_by('-analysis_time')
        data = [
            {
                'id': item.id,
                'file_name': item.file_name,
                'audit_results': item.audit_results,
                'analysis_time': item.analysis_time.isoformat()
            }
            for item in history
        ]
        return JsonResponse(data, safe=False)



def get_analysis_by_id(request, session_id):
    try:
        obj = ContractAnalysis.objects.get(id=session_id)
        return JsonResponse({
            'id': obj.id,
            'file_name': obj.file_name,
            'audit_results': obj.audit_results,
            'pdf_file': request.build_absolute_uri(obj.pdf_file.url),  # 这里加上完整文件URL
        })
    except ContractAnalysis.DoesNotExist:
        return JsonResponse({'error': 'Not found'}, status=404)