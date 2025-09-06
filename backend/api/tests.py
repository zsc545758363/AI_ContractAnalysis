from django.test import TestCase
import pythoncom
import win32com.client as win32
import os
from docx import Document


def extract_docx_text(file_path):
    """使用python-docx从docx文件中提取纯文本内容"""
    try:
        doc = Document(file_path)
        full_text = []

        # 遍历文档中的每个段落
        for para in doc.paragraphs:
            full_text.append(para.text)

        # 遍历文档中的每个表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        # 将所有文本连接成一个字符串
        return '\n'.join(full_text)
    except Exception as e:
        print(f"提取docx文本失败: {e}")
        return None


def convert_doc_to_docx_and_extract(file_path):
    """将doc文件转换为docx并提取文本"""
    # 确保路径正确
    safe_file_path = os.path.abspath(file_path)

    # 初始化COM环境
    pythoncom.CoInitialize()
    try:
        # 创建Word应用实例
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False

        # 转换文件
        docx_path = safe_file_path + ".converted.docx"
        doc = word.Documents.Open(safe_file_path)
        doc.SaveAs(docx_path, FileFormat=16)  # 16 = wdFormatDocumentDefault (.docx)
        doc.Close()

        # 提取文本
        text = extract_docx_text(docx_path)

        # 清理临时文件
        if os.path.exists(docx_path):
            os.remove(docx_path)

        return text
    except Exception as e:
        print(f"转换doc文件失败: {e}")
        return None
    finally:
        # 关闭Word应用并释放COM资源
        word.Quit()
        pythoncom.CoUninitialize()


from openai import OpenAI


# 输出清洗
def parse_json_with_code_blocks(json_text):
    """
    解析带Markdown代码块标记的JSON字符串
    支持处理 ```json 和 ``` 包裹的JSON数据
    """
    try:
        # 去除开头的 ```json 或 ```
        if json_text.startswith("```json"):
            json_text = json_text[len("```json"):]
        elif json_text.startswith("```"):
            json_text = json_text[len("```"):]

        # 去除末尾的 ```
        if json_text.endswith("```"):
            json_text = json_text[:-len("```")]

        # 去除前后可能存在的空白字符
        json_text = json_text.strip()

        # 解析JSON
        data = json.loads(json_text)
        return data
    except json.JSONDecodeError as e:
        print(f"JSON解析失败: {e}")
        print(f"待解析文本: {json_text[:100]}...")  # 打印前100个字符用于调试
        return None
    except Exception as e:
        print(f"处理过程中发生错误: {e}")
        return None

import json


# 使用示例
if __name__ == "__main__":
    # 使用原始字符串或双反斜杠表示Windows路径
    docx_path = r"C:\Python project\ContractAnalysis\Demo_Source\景津装备压滤机滤板购销合同25.5.8.docx"
    #doc_path = r"C:\Python project\ContractAnalysis\Demo_Source\合同审核要点.doc"

    # 测试docx文件提取
    text = extract_docx_text(docx_path)
    if text:
        print("从docx文件提取的文本:")
        print(text)  # 打印前200个字符

    # # 测试doc文件转换和提取
    # doc_text = convert_doc_to_docx_and_extract(doc_path)
    # if doc_text:
    #     print("从doc文件提取的文本:")
    #     print(type(doc_text))  # 打印前200个字符

    prompt = f"""
    请你作为合同审核专家，从以下9个方面分析以零星采购合同内容是否符合公司合规要求：
    1.供方资质审核：是否存在失信情况，是否有重大负债（100万元以上）；
    2.合同编号中日期与签订时间一致；
    3.产品名称前后一致，合同总金额计算无误，大小写一致；
    4.交货时间和地点清晰明确；
    5.结算方式为承兑付款，付款时间清晰明确；
    6.违约责任包含逾期到货违约金和逾期到货合同解除权；
    7.争议解决的诉讼法院为需方所在地或原告方所在法院；
    8.签订时间已明确，无空白条款。
    9.该合同是否与《民法典》及公司规章制度相冲突。
    对于每一个点，你需要返回三个字段：
    "status": "✅ 符合 /⚠️ 部分风险 / ❌ 严重风险",
    "detail": "该点在合同原文中的体现",
    "risk": "如存在风险，简要描述风险，否则填-"
    你需要严格按照这个格式返回数据，最后将这个九个点的分析结果按顺序以JSON格式返回，示例输出格式如下：
    [
      {{
        "status": "⚠️ 部分风险",
        "detail": "合同中未涉及供方（景津装备股份有限公司）的失信记录或负债情况。需通过外部渠道（如国家企业信用信息公示系统、天眼查等）独立核实其是否存在失信、重大负债（>100万元）等问题。",
        "risk": "若供方存在失信/大额负债，可能影响履约能力"
      }},
      {{
        "status": "✅ 符合",
        "detail": "合同编号：JAJJ20250508504（编号含日期20250508） 签订日期：2025年5月8日",
        "risk": "-"
      }},
      {{
        "status": "✅ 符合",
        "detail": "产品名称一致性：
    压滤机压紧板（规格1250）

    压滤机止推板（规格1250）

    压滤机滤板（规格1米*1米）
    ✅ 名称前后无矛盾。

    金额计算：
    产品	单价	数量	小计
    压紧板	16,000	1	16,000
    止推板	19,000	1	19,000
    滤板	750	60	45,000
    合计（含税）			80,000元
    大写金额：捌万圆整

    小写金额：￥：80,000.00元
    ✅ 总金额计算正确，大小写一致。",
        "risk": "-"
      }},
      ...(剩余6个点以此类推)
    ]

    需审核的合同如下：
    {text}

    请你将分析结果按格式返回。
    """




    client = OpenAI(api_key="sk-ca6282f72dcd42798d469ea66ffc02a2", base_url="https://api.deepseek.com")

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": "你是一个专业的企业合同审核助手"},
            {"role": "user", "content": prompt},
        ],
        stream=False
    )

    print(response.choices[0].message.content)
    result_text = response.choices[0].message.content
    result_text = parse_json_with_code_blocks(result_text)

    print(result_text)

